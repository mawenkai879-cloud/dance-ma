"""Generate a Word (.docx) containing ONLY the Method section, following the
formatting conventions of the EDGE paper (CVPR 2023): a numbered "3. Method"
section with subsections (Background, Architecture / hierarchical conditioning,
Training losses, Editing-style global constraint).
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor


def set_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    # body paragraphs justified, small first-line indent like a paper
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0


def add_section_heading(doc, text, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(12)
    return p


def add_equation(doc, text):
    """Centered, italic equation line (plain-text math, paper-like)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p


def main():
    doc = Document()
    set_base_styles(doc)

    # ---- 3. Method ----
    add_section_heading(doc, "3. Method", size=13)
    add_body(doc,
        "We build on a transformer-based dance diffusion model that generates a "
        "motion sequence conditioned on per-frame music features. Our key "
        "observation is that conditioning each frame solely on its local music "
        "feature provides no explicit representation of the global style of a "
        "clip. We therefore introduce a two-level hierarchy: a global branch "
        "that summarizes the entire clip into a style vector, and a local branch "
        "in which this global style guides per-frame generation. We further "
        "propose a hierarchical training loss that emphasizes chosen time "
        "segments and body parts. We first review the diffusion background "
        "(Sec. 3.1), then describe our hierarchical conditioning (Sec. 3.2), the "
        "hierarchical loss (Sec. 3.3), and finally the sampling-time global "
        "constraint (Sec. 3.4).")

    # 3.1 Background
    add_section_heading(doc, "3.1. Background: Diffusion for motion")
    add_body(doc,
        "We adopt a Denoising Diffusion Probabilistic Model. Given a clean motion "
        "x0, the forward process gradually adds Gaussian noise:")
    add_equation(doc,
        "q(x_t | x_0) = N( x_t ; sqrt(a_bar_t) x_0 , (1 - a_bar_t) I ),")
    add_equation(doc,
        "x_t = sqrt(a_bar_t) x_0 + sqrt(1 - a_bar_t) eps ,   eps ~ N(0, I),")
    add_body(doc,
        "where a_bar_t is the cumulative product of the noise schedule. A network "
        "f_theta(x_t, c, t) is trained to denoise x_t given the music condition c "
        "and timestep t, predicting the clean signal x0. The motion tensor has "
        "T = 150 frames (5 s at 30 fps) and D = 151 channels: 4 foot-contact, 3 "
        "root-translation, and 24x6 joint rotations in the 6D rotation "
        "representation; joint j occupies rotation channels [7+6j, 13+6j). The "
        "music feature c has F = 4800 channels per frame (Jukebox). Our "
        "contributions modify only how the condition enters f_theta and how the "
        "loss is weighted; the diffusion schedule is unchanged.")

    # 3.2 Hierarchical conditioning
    add_section_heading(doc, "3.2. Hierarchical conditioning: global to local")
    add_subheading(doc, "Global style vector.")
    add_body(doc,
        "We mean-pool the music feature over time and encode it with an MLP to "
        "obtain a global style vector g, the first hierarchical feature, which "
        "captures the overall character of the clip independently of any single "
        "frame:")
    add_equation(doc,
        "c_bar = (1/T) sum_tau c_tau ,    g = MLP_global( c_bar ).")

    add_subheading(doc, "Global modulates local (FiLM).")
    add_body(doc,
        "The global style modulates the per-frame local condition tokens h_loc "
        "through a feature-wise affine (FiLM) transform, so that the same local "
        "beats are re-expressed under the chosen global style:")
    add_equation(doc,
        "(gamma, beta) = DenseFiLM(g),   h_loc_tau' = (1 + gamma) * h_loc_tau + beta.")

    add_subheading(doc, "Dual-channel injection.")
    add_body(doc,
        "Beyond modulating the local tokens, the global token is injected in two "
        "additional places so the decoder can attend to it directly: (a) it is "
        "appended to the cross-attention memory C = [ h_loc' ; h_time ; g ]; and "
        "(b) it is added to the FiLM timestep condition, t_cond <- t_cond + g. The "
        "transformer decoder then denoises x_t using C as memory and t_cond as the "
        "FiLM condition.")

    add_subheading(doc, "Classifier-free guidance.")
    add_body(doc,
        "We introduce a learnable null global token g_0 (null_global_token). "
        "During training the condition is dropped with probability p_drop, in "
        "which case g is replaced by g_0, synchronized with the local-condition "
        "dropout. At sampling time we mix the conditional and unconditional "
        "predictions with guidance weight w:")
    add_equation(doc,
        "f_guided = f_uncond + w ( f_cond - f_uncond ).")

    # 3.3 Hierarchical loss
    add_section_heading(doc, "3.3. Hierarchical loss: time-segment x body-part")
    add_body(doc,
        "Given a configuration with a time segment [a, b), a set of joints J, and "
        "an emphasis factor w, we build multiplicative weight fields that equal w "
        "inside the chosen time segment AND the chosen body parts, and 1 "
        "elsewhere. Let w_e = w - 1, M_time(tau) = 1[a <= tau < b], and "
        "M_joint(j) = 1[j in J]:")
    add_equation(doc,
        "W_chan(tau, k) = 1 + w_e * M_time(tau) * M_chan(k),")
    add_equation(doc,
        "W_joint(tau, j) = 1 + w_e * M_time(tau) * M_joint(j).")
    add_body(doc,
        "These fields weight the reconstruction loss (channel level) and the "
        "forward-kinematics loss (joint level), letting the model allocate more "
        "fidelity to user-chosen frames and joints. The total objective keeps the "
        "four-term weighting of the backbone:")
    add_equation(doc,
        "L = 0.636 L_rec + 2.964 L_vel + 0.646 L_fk + lambda_foot L_foot,")
    add_equation(doc,
        "L_rec = || W_chan * ( f_theta(x_t,c,t) - x_0 ) ||,")
    add_equation(doc,
        "L_fk  = || W_joint * ( FK(x0_hat) - FK(x0) ) ||,")
    add_body(doc,
        "where L_vel is a first-order velocity loss, L_foot penalizes foot "
        "velocity on frames the model itself predicts as in-contact "
        "(lambda_foot = 10.942), and FK denotes forward kinematics mapping joint "
        "rotations to 3D joint positions. For joint j the rotation channels "
        "[7+6j, 13+6j) are selected; selecting the root joint additionally "
        "includes its translation and contact channels [0, 7).")

    # 3.4 Editing-style global constraint
    add_section_heading(doc, "3.4. Sampling-time global constraint")
    add_body(doc,
        "Prior editing approaches (inpainting with binary masks) impose hard, "
        "local constraints at sampling time: chosen joints or frames are forced to "
        "given values while the rest is inpainted. This requires the target values "
        "up front and has no notion of global style. Our hierarchy adds a "
        "complementary, global control axis. By overriding the pooled feature with "
        "a reference clip's pooled feature, c_bar <- c_bar_ref, the global style is "
        "swapped while the local beats are preserved:")
    add_equation(doc,
        "g = MLP_global( c_bar_ref ),   ( c_bar_ref from another clip / target style ).")
    add_body(doc,
        "This first-level (global) constraint cannot be expressed by local "
        "inpainting, and requires no target motion - only a style vector. Together, "
        "hierarchical conditioning provides a global control that local editing "
        "lacks, while the hierarchical loss acts as a training prior rather than a "
        "runtime hard mask.")

    out = "docs/Method.docx"
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
