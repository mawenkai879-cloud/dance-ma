"""Generate the detailed Chinese Method write-up (docs/方法详述.md) as a Word
(.docx) file: docs/方法详述.docx. Uses SimSun for CJK body text and adds
centered plain-text equations and inline code blocks."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_cjk(run, font="宋体"):
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def base(doc):
    n = doc.styles["Normal"]
    n.font.size = Pt(10.5)
    n.paragraph_format.line_spacing = 1.15
    n.paragraph_format.space_after = Pt(4)
    n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    set_cjk(r, "黑体")
    p.paragraph_format.space_after = Pt(10)


def h(doc, text, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    set_cjk(r, "黑体")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def sub(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    set_cjk(r, "黑体")
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cjk(r, "宋体")
    p.paragraph_format.first_line_indent = Pt(21)
    return p


def eq(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def code(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        set_cjk(run, "黑体")
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            set_cjk(run, "宋体")
            run.font.size = Pt(9)


def main():
    doc = Document()
    base(doc)

    title(doc, "方法（详细版）")
    body(doc,
        "本文档详细阐述本工作提出的层次化音乐驱动舞蹈生成方法。整体框架建立在基于 "
        "Transformer 的扩散舞蹈生成模型之上，核心贡献是在原有“逐帧局部条件”之外，"
        "显式引入全局（整曲风格）到局部（逐帧细节）的层次化结构，并配以一个时间段 × "
        "身体部位的层次化损失。")

    # 1
    h(doc, "1  问题定义与符号")
    body(doc,
        "给定一段音乐特征序列，模型学习生成对应的 3D 人体舞蹈动作。记动作张量 "
        "x ∈ R^(B×T×D)，其中 B 为 batch，T=150 帧（30 fps，5 秒），D=151 为动作通道；"
        "音乐特征 c ∈ R^(B×T×F)，为逐帧的 Jukebox 特征，F=4800；扩散时间步 "
        "t ∈ {1,…,N}，N=1000。动作通道 D=151 的布局为：4 维脚触地 + 3 维根平移 + "
        "24 关节×6D 旋转；第 j 个关节的旋转占据通道 [7+6j, 13+6j)。")
    body(doc,
        "基线的局限：原模型每一帧 x_τ 仅以该帧的局部音乐特征 c_τ 为条件，没有任何对"
        "整段音乐风格的显式表征，导致全局编排的一致性完全依赖网络隐式学习。")

    # 2
    h(doc, "2  扩散框架（预备）")
    body(doc,
        "采用 DDPM 前向加噪过程，对干净动作 x0：")
    eq(doc, "q(x_t | x_0) = N( x_t ; √(ā_t)·x_0 , (1−ā_t)·I ),")
    eq(doc, "x_t = √(ā_t)·x_0 + √(1−ā_t)·ε ,   ε ~ N(0, I),")
    body(doc,
        "其中 ā_t = ∏_{s≤t}(1−β_s)。去噪网络 f_θ(x_t, c, t) 直接预测干净信号 x0。"
        "本方法的全部改动都集中在条件如何进入 f_θ 以及损失如何加权，扩散调度本身不变。")

    # 3
    h(doc, "3  层次化条件：全局 → 局部")
    body(doc, "实现见 model/model.py 的 DanceDecoder.forward。")

    sub(doc, "3.1  全局风格向量（第一层特性）")
    body(doc,
        "对整段音乐在时间维做均值池化，再经 MLP 编码，得到全局风格向量 g ∈ R^(B×d)："
        "它刻画整曲的风格/能量，与任意单帧无关，是层次结构的第一层特性。")
    eq(doc, "c̄ = (1/T)·Σ_τ c_τ ,    g = MLP_global( c̄ ).")
    code(doc,
        "global_pool   = cond_embed.mean(dim=1)        # (B, F)\n"
        "global_hidden = self.global_encoder(global_pool)  # g: (B, d)")

    sub(doc, "3.2  全局调制局部（FiLM）")
    body(doc,
        "将逐帧的局部条件 token h_loc 用 g 做特征级仿射调制（FiLM）：由 g 生成缩放 γ "
        "与平移 β，对每一帧施加。同样的局部鼓点，在不同全局风格下被重新表达。")
    eq(doc, "(γ, β) = DenseFiLM(g),   h_loc_τ′ = (1+γ)⊙h_loc_τ + β.")

    sub(doc, "3.3  双通道注入")
    body(doc,
        "除调制局部 token 外，全局 token 还被注入解码器的另外两处：(a) 拼接到交叉注意力"
        "记忆序列尾部 C = [ h_loc′ ; h_time ; g ]；(b) 加到 FiLM 时间步条件 "
        "t_cond ← t_cond + g。随后 Transformer 解码器以 C 为记忆、t_cond 为 FiLM 条件，"
        "对带噪动作 x_t 去噪。")
    code(doc,
        "c = torch.cat((cond_tokens, t_tokens), dim=-2)\n"
        "global_token = global_hidden.unsqueeze(1)\n"
        "global_token = torch.where(keep_mask_embed, global_token, null_global_token)\n"
        "c = torch.cat((c, global_token), dim=-2)                 # (a) 交叉注意力记忆\n"
        "t = t + torch.where(keep_mask_hidden, global_hidden, 0)  # (b) FiLM 时间条件")

    sub(doc, "3.4  无分类器引导（CFG）")
    body(doc,
        "引入可学习的空全局 token g∅（null_global_token）。训练时以概率 p_drop 丢弃条件，"
        "此时用 g∅ 替换 g（与局部条件丢弃同步）。采样时按引导权重 w 混合有/无条件预测：")
    eq(doc, "f_guided = f_∅ + w·( f_c − f_∅ ).")

    sub(doc, "3.5  采样期全局约束 global_pool_override")
    body(doc,
        "提供一个钩子：采样时直接指定全局风格，而非从音乐计算。这正是局部 inpainting "
        "无法表达的第一层（全局）约束。把 global_pool_override 设为另一首曲子的池化特征，"
        "即可在保持局部鼓点不变的前提下替换全局风格——这是风格交换实验的基础。")
    eq(doc, "c̄ ← c̄_ref ,    g = MLP_global( c̄_ref ).")

    # 4
    h(doc, "4  层次化损失：时间段 × 身体部位")
    body(doc, "实现见 model/diffusion.py 的 _hierarchical_weights 与 p_losses。")
    sub(doc, "4.1  权重场")
    body(doc,
        "给定配置 {frames=[a,b), joints=J, weight=w}，构造两个乘性权重场，除选中的时间段"
        "且选中的身体部位取值 w 外，其余处处为 1。令 w_e=w−1，M_time(τ)=1[a≤τ<b]，"
        "M_joint(j)=1[j∈J]：")
    eq(doc, "W_chan(τ,k) = 1 + w_e·M_time(τ)·M_chan(k),")
    eq(doc, "W_joint(τ,j) = 1 + w_e·M_time(τ)·M_joint(j).")
    body(doc,
        "其中关节 j 对应通道 [7+6j, 13+6j)；若选中根关节 0，其平移与触地通道 [0,7) 也一并选入。")

    sub(doc, "4.2  加权到重建与 FK 损失")
    body(doc,
        "总损失沿用基线的四项加权和，本方法仅对其中重建项与 FK 项施加上述权重场：")
    eq(doc, "L = 0.636·L_rec + 2.964·L_vel + 0.646·L_fk + λ_foot·L_foot,")
    eq(doc, "L_rec = ‖ W_chan ⊙ ( f_θ(x_t,c,t) − x_0 ) ‖,")
    eq(doc, "L_fk  = ‖ W_joint ⊙ ( FK(x̂_0) − FK(x_0) ) ‖,")
    body(doc,
        "其中 L_vel 为速度（一阶差分）损失，L_foot 为脚滑损失（仅在模型自身预测触地的帧上"
        "惩罚脚部速度），λ_foot=10.942；FK 表示前向运动学，将关节旋转映射到 3D 关节位置。")

    sub(doc, "4.3  命令行接口")
    code(doc,
        "--hierarchical_loss_frames 0,75            # 时间段 [0, 75)\n"
        "--hierarchical_loss_joints 16,17,18,19,20  # SMPL 关节索引（肩/肘/腕）\n"
        "--hierarchical_loss_weight 1.5             # 强调系数 w")
    body(doc,
        "SMPL 关节索引（见 dataset/masks.py）：0 根，16/17 肩，18/19 肘，20/21 腕，"
        "下肢 1,2,4,5,7,8,10,11。")

    # 5
    h(doc, "5  与 EDGE 编辑能力的关系（非重复性论证）")
    body(doc,
        "EDGE 的编辑（inpaint_loop + dataset/masks.py 掩码）是采样期的局部硬约束：指定"
        "关节/帧被强制设为给定真值，其余 inpaint 补全。它需要事先提供目标值，且没有"
        "“全局风格”的概念。本方法的两个模块与之层级不同、互补：")
    table(doc,
        ["维度", "EDGE inpainting", "层次化条件（本文）", "层次化损失（本文）"],
        [
            ["作用层级", "局部（关节/帧）", "全局（整曲风格）", "训练期区域强调"],
            ["作用时机", "采样期", "采样期（override）", "训练期"],
            ["约束类型", "硬替换", "风格引导", "软保真度加权"],
            ["需目标动作", "是", "否（只需风格向量）", "否"],
        ])
    body(doc,
        "即：层次化条件提供了 EDGE 没有的全局控制轴；层次化损失是训练先验而非运行时硬掩码。")

    # 6
    h(doc, "6  评测协议")
    body(doc,
        "数据：AIST++ 测试集，20 段去重切片；同一变体间使用相同采样噪声以保证公平。"
        "指标：PFC（脚滑，越低越好）；Beat Align（节拍对齐，越高越好）；FID_k（相对真值"
        "分布的 Fréchet 距离，越低越好）；Div_k（动力学多样性，越接近真值 10.21 越好）。")
    sub(doc, "受控消融主结果（统一 10 epoch 微调、损失权重 1.5）")
    table(doc,
        ["变体", "全局条件", "层次损失", "PFC↓", "Beat↑", "FID_k↓", "Div_k(→10.21)"],
        [
            ["基线", "—", "—", "1.019", "0.239", "341.6", "13.82"],
            ["+层次化条件", "✓", "—", "1.736", "0.257", "423.3", "19.04"],
            ["+层次化损失", "—", "✓", "1.323", "0.250", "427.5", "17.11"],
            ["+两者（完整）", "✓", "✓", "0.619", "0.264", "170.6", "4.91"],
        ])
    body(doc,
        "完整方法在 PFC、Beat Align、FID_k 上均最优：两个模块单独使用时不稳定，但组合后"
        "取得最低脚滑、最佳节拍对齐，FID 约为基线的一半，体现出明确的协同效应。")

    out = "docs/方法详述.docx"
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
